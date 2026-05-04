"""
Gerador Automático de Relatório de Atendimento Telefónico
- Lê ficheiro Excel de dados de chamadas (geral + opcional manhã/tarde)
- Usa template Word com header/footer institucionais
- Gera relatório profissional Word com análise por IA
"""

import streamlit as st
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import io, re, json, os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL

import anthropic
try:
    from config import ANTHROPIC_API_KEY as _CFG_KEY, ANTHROPIC_MODEL as _CFG_MODEL, MAX_TOKENS as _CFG_TOKENS
except ImportError:
    _CFG_KEY = ""; _CFG_MODEL = "claude-opus-4-5"; _CFG_TOKENS = 3200

# ──────────────────────────────────────────────
#  MAPPING: GeralMXXX → assunto legível
# ──────────────────────────────────────────────
def load_mapping(path=None):
    candidates = [
        path,
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "grupos_mapping.json"),
        "grupos_mapping.json",
    ]
    for p in candidates:
        if p and os.path.exists(p):
            with open(p, encoding="utf-8") as f:
                return json.load(f)
    return {}

_MAPPING = load_mapping()

def label(group: str, short=False) -> str:
    info = _MAPPING.get(group, {})
    if not info:
        return group
    key = "assunto_curto" if short else "assunto"
    return info.get(key) or info.get("assunto") or group

def area(group: str) -> str:
    return _MAPPING.get(group, {}).get("area", "")

def servico(group: str) -> str:
    return _MAPPING.get(group, {}).get("servico", "")

# ──────────────────────────────────────────────
#  PATHS
# ──────────────────────────────────────────────
_DIR = os.path.dirname(os.path.abspath(__file__))
HEADER_IMG = os.path.join(_DIR, "template_header.png")
FOOTER_IMG = os.path.join(_DIR, "template_footer.png")

# ──────────────────────────────────────────────
#  COLOUR PALETTE
# ──────────────────────────────────────────────
C_NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
C_BLUE   = RGBColor(0x4A, 0x90, 0xD9)
C_STEEL  = RGBColor(0x7F, 0xB3, 0xD3)
C_TEAL   = RGBColor(0x2E, 0x86, 0xAB)
C_RED    = RGBColor(0xE8, 0x48, 0x55)
C_GREEN  = RGBColor(0x27, 0xAE, 0x60)
C_ORANGE = RGBColor(0xF3, 0x9C, 0x12)
C_TEXT   = RGBColor(0x2C, 0x3E, 0x50)
C_GREY   = RGBColor(0x7F, 0x8C, 0x8D)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_LBLUE  = RGBColor(0xF0, 0xF5, 0xFA)

HEX_NAVY  = "1B3A5C"
HEX_BLUE  = "4A90D9"
HEX_TEAL  = "2E86AB"
HEX_RED   = "E84855"
HEX_GREEN = "27AE60"
HEX_ORANGE= "F39C12"
HEX_LGREY = "F0F5FA"
HEX_WHITE = "FFFFFF"

# ──────────────────────────────────────────────
#  MERGE PAIRS  (Ajuda de Assuntos + homónimos)
# ──────────────────────────────────────────────
MERGE_PAIRS = [
    # Ajuda de Assuntos homónimos
    {"canonical":"GeralM212","groups":["GeralM212","GeralM501"],
     "assunto":"Benefício Fiscal ao Gasóleo Agrícola","assunto_curto":"Benef. Fiscal Gasóleo"},
    {"canonical":"GeralM213","groups":["GeralM213","GeralM502"],
     "assunto":"Programa Emparcelar para Ordenar","assunto_curto":"Emparcelar para Ordenar"},
    {"canonical":"GeralM214","groups":["GeralM214","GeralM503"],
     "assunto":"Homologação de Tratores Agrícolas / Mecanização Agrária","assunto_curto":"Homologação Tratores"},
    {"canonical":"GeralM311","groups":["GeralM311","GeralM506"],
     "assunto":"Rede Rural PAC","assunto_curto":"Rede Rural PAC"},
    {"canonical":"GeralM321","groups":["GeralM321","GeralM507"],
     "assunto":"RAN — Reserva Agrícola Nacional","assunto_curto":"RAN"},
    {"canonical":"GeralM234","groups":["GeralM234","GeralM508"],
     "assunto":"Bem-Estar Animal PEPAC","assunto_curto":"Bem-Estar Animal"},
    # a) Portaria: Linha de Apoio + Geral + Apoio às Explorações Agrícolas
    {"canonical":"GeralM509","groups":["GeralM509","GeralM909","GeralM219"],
     "assunto":"Atendimento por Operador — Portaria",
     "assunto_curto":"Atendimento por Operador — Portaria"},
    # d) DSTAR: Ordenamento + Estruturas Agro-Rurais (mesma divisão)
    {"canonical":"GeralM329","groups":["GeralM329","GeralM339"],
     "assunto":"Atendimento por Operador — Secretariado da DSTAR — Divisão de Ordenamento do Espaço Rural",
     "assunto_curto":"Secretariado DSTAR — Ordenamento"},
]
_MERGE_LOOKUP = {}
for _mp in MERGE_PAIRS:
    for _g in _mp["groups"]:
        _MERGE_LOOKUP[_g] = _mp


def _hms_to_sec(s: str) -> int:
    try:
        parts = [int(x) for x in str(s).split(":")]
        if len(parts) == 3:
            return parts[0]*3600 + parts[1]*60 + parts[2]
        if len(parts) == 2:
            return parts[0]*60 + parts[1]
    except Exception:
        pass
    return 0

def _sec_to_hms(total: int) -> str:
    h = total // 3600; m = (total % 3600)//60; s = total % 60
    return f"{h}:{m:02d}:{s:02d}"


def consolidate_groups(groups: list) -> list:
    merged_buckets: dict = {}
    standalone: list = []
    for g in groups:
        mp = _MERGE_LOOKUP.get(g["group"])
        if mp is None:
            standalone.append(g)
        else:
            merged_buckets.setdefault(mp["canonical"], []).append(g)
    result = list(standalone)
    for canonical_code, bucket in merged_buckets.items():
        mp = _MERGE_LOOKUP[canonical_code]
        presented = sum(g["presented"] for g in bucket)
        answered  = sum(g["answered"]  for g in bucket)
        missed    = sum(g["missed"]    for g in bucket)
        tot_sec   = sum(_hms_to_sec(g["total_talking"]) for g in bucket)
        tot_tlk   = _sec_to_hms(tot_sec)
        avg_tlk   = _sec_to_hms(tot_sec // answered) if answered > 0 else "N/D"
        pct_ans   = answered / presented if presented > 0 else 0.0
        priority  = round(missed * (1 - pct_ans), 1)
        canon_g   = next((g for g in bucket if g["group"] == canonical_code), bucket[0])
        result.append({
            "group": canonical_code,
            "label": mp["assunto"],
            "label_short": mp["assunto_curto"],
            "area": canon_g.get("area", ""),
            "servico": canon_g.get("servico", ""),
            "presented": presented, "answered": answered, "missed": missed,
            "pct_answered": round(pct_ans, 4),
            "response_rate": round(pct_ans*100, 1),
            "avg_speed": canon_g.get("avg_speed","N/D"),
            "total_talking": tot_tlk, "avg_talking": avg_tlk,
            "priority_index": priority,
            "merged_from": [g["group"] for g in bucket],
        })
    result.sort(key=lambda x: x["priority_index"], reverse=True)
    return result


# ──────────────────────────────────────────────
#  EXCEL PARSING — ficheiro principal
# ──────────────────────────────────────────────
def parse_excel(file):
    summary = {}
    groups  = []

    df_s = pd.read_excel(file, sheet_name="Summary", header=None)
    row_dates = df_s.iloc[2]
    row_data  = df_s.iloc[4]
    try:
        start_date = pd.to_datetime(row_dates[0]).strftime("%d/%m/%Y")
        end_date   = pd.to_datetime(row_dates[1]).strftime("%d/%m/%Y")
    except Exception:
        start_date = str(row_dates[0])[:10]
        end_date   = str(row_dates[1])[:10]

    summary = {
        "start_date": start_date, "end_date": end_date,
        "total_presented": int(row_data[0]) if pd.notna(row_data[0]) else 0,
        "total_answered":  int(row_data[1]) if pd.notna(row_data[1]) else 0,
        "total_missed":    int(row_data[2]) if pd.notna(row_data[2]) else 0,
        "pct_answered":   float(row_data[5]) if pd.notna(row_data[5]) else 0,
        "pct_missed":     float(row_data[6]) if pd.notna(row_data[6]) else 0,
        "total_talking":  str(row_data[7]) if pd.notna(row_data[7]) else "N/D",
        "avg_talking":    str(row_data[8]) if pd.notna(row_data[8]) else "N/D",
    }

    df_d = pd.read_excel(file, sheet_name="NO SUB GROUP", header=None)
    for i in range(3, len(df_d)):
        row = df_d.iloc[i]
        grp = str(row[0]).strip() if pd.notna(row[0]) else ""
        if not grp or grp in ("nan",""): continue
        presented = int(row[1]) if pd.notna(row[1]) else 0
        answered  = int(row[2]) if pd.notna(row[2]) else 0
        missed    = int(row[3]) if pd.notna(row[3]) else 0
        if presented == 0: continue
        pct_ans = float(row[6]) if pd.notna(row[6]) else 0.0
        priority = round(missed * (1 - pct_ans), 1)
        groups.append({
            "group": grp, "label": label(grp),
            "label_short": label(grp, short=True),
            "area": area(grp), "servico": servico(grp),
            "presented": presented, "answered": answered, "missed": missed,
            "pct_answered": pct_ans, "response_rate": round(pct_ans*100,1),
            "avg_speed": str(row[8]) if pd.notna(row[8]) else "N/D",
            "total_talking": str(row[9]) if pd.notna(row[9]) else "N/D",
            "avg_talking": str(row[10]) if pd.notna(row[10]) else "N/D",
            "priority_index": priority,
        })
    groups = consolidate_groups(groups)
    return summary, groups


# ──────────────────────────────────────────────
#  EXCEL PARSING — ficheiros de período (manhã/tarde)
# ──────────────────────────────────────────────
def parse_excel_period(file, period_label: str) -> dict:
    """Return dict {group_code: {presented, answered, missed, response_rate, label}}"""
    df = pd.read_excel(file, sheet_name="NO SUB GROUP", header=None)
    groups = {}
    for i in range(3, len(df)):
        row = df.iloc[i]
        grp = str(row[0]).strip() if pd.notna(row[0]) else ""
        if not grp or grp in ("nan",""): continue
        if not re.match(r'^GeralM\d+$', grp): continue
        presented = int(float(row[1])) if pd.notna(row[1]) else 0
        answered  = int(float(row[2])) if pd.notna(row[2]) else 0
        missed    = int(float(row[3])) if pd.notna(row[3]) else 0
        if presented == 0: continue
        pct = float(row[6]) if pd.notna(row[6]) else 0.0
        groups[grp] = {
            "group": grp, "label": label(grp), "label_short": label(grp, short=True),
            "presented": presented, "answered": answered, "missed": missed,
            "response_rate": round(pct*100, 1), "period": period_label,
        }
    # Merge homonyms within the period
    merged_buckets: dict = {}
    standalone: dict = {}
    for g in groups.values():
        mp = _MERGE_LOOKUP.get(g["group"])
        if mp is None:
            standalone[g["group"]] = g
        else:
            merged_buckets.setdefault(mp["canonical"], []).append(g)
    result = dict(standalone)
    for canonical_code, bucket in merged_buckets.items():
        mp = _MERGE_LOOKUP[canonical_code]
        presented = sum(g["presented"] for g in bucket)
        answered  = sum(g["answered"]  for g in bucket)
        missed    = sum(g["missed"]    for g in bucket)
        pct       = answered/presented if presented > 0 else 0.0
        result[canonical_code] = {
            "group": canonical_code, "label": mp["assunto"],
            "label_short": mp["assunto_curto"],
            "presented": presented, "answered": answered, "missed": missed,
            "response_rate": round(pct*100, 1), "period": period_label,
        }
    return result


# ──────────────────────────────────────────────
#  CHART GENERATION
# ──────────────────────────────────────────────
NAVY_MPL  = "#1B3A5C"
BLUE_MPL  = "#4A90D9"
TEAL_MPL  = "#2E86AB"
RED_MPL   = "#E84855"
GREEN_MPL = "#27AE60"
AMBER_MPL = "#F39C12"

def _fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return buf.getvalue()


def make_charts(summary, groups):
    charts = []

    # 1 ── Pie: answered vs missed ──
    fig, ax = plt.subplots(figsize=(5,5), facecolor="white")
    sizes  = [summary["total_answered"], summary["total_missed"]]
    colors = [TEAL_MPL, RED_MPL]
    lbls   = [f"Atendidas\n{summary['total_answered']}",
              f"Perdidas\n{summary['total_missed']}"]
    wedges, texts, autotexts = ax.pie(
        sizes, labels=lbls, colors=colors, autopct="%1.1f%%",
        startangle=90, pctdistance=0.78,
        textprops={"fontsize":11,"color":"#2C3E50"},
        wedgeprops={"linewidth":2.5,"edgecolor":"white"},
    )
    for at in autotexts:
        at.set_color("white"); at.set_fontweight("bold"); at.set_fontsize(11)
    ax.set_title("Distribuição Global de Chamadas", fontsize=13,
                 fontweight="bold", color=NAVY_MPL, pad=14)
    charts.append(_fig_to_bytes(fig))

    # 2 ── Bar: top-15 by volume ──
    top15 = sorted(groups, key=lambda x: x["presented"], reverse=True)[:15]
    names = [g["label_short"] for g in top15]
    ans   = [g["answered"] for g in top15]
    msd   = [g["missed"]   for g in top15]
    x = np.arange(len(names)); w = 0.38
    fig, ax = plt.subplots(figsize=(12,5.5), facecolor="white")
    ax.bar(x-w/2, ans, w, label="Atendidas", color=TEAL_MPL, linewidth=0)
    ax.bar(x+w/2, msd, w, label="Perdidas",  color=RED_MPL,  linewidth=0)
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=40, ha="right", fontsize=8.5)
    ax.set_ylabel("Nº de Chamadas", fontsize=11, color="#2C3E50")
    ax.set_title("Top 15 Grupos — Atendidas vs. Perdidas", fontsize=13,
                 fontweight="bold", color=NAVY_MPL, pad=14)
    ax.legend(fontsize=10, framealpha=0)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.set_facecolor("white"); ax.yaxis.grid(True, linestyle="--", alpha=0.4)
    ax.set_axisbelow(True)
    charts.append(_fig_to_bytes(fig))

    # 3 ── Horizontal bar: top-10 priority index ──
    top10p  = sorted(groups, key=lambda x: x["priority_index"], reverse=True)[:10]
    pnames  = [g["label_short"] for g in top10p]
    pvals   = [g["priority_index"] for g in top10p]
    cmap    = plt.cm.RdYlBu_r
    norm    = plt.Normalize(min(pvals) if pvals else 0, max(pvals) if pvals else 1)
    bcolors = [cmap(norm(v)) for v in pvals]
    fig, ax = plt.subplots(figsize=(9,5), facecolor="white")
    bars = ax.barh(pnames[::-1], pvals[::-1], color=bcolors[::-1], height=0.6)
    ax.set_xlabel("Índice de Prioridade", fontsize=11, color="#2C3E50")
    ax.set_title("Top 10 Grupos por Índice de Prioridade", fontsize=13,
                 fontweight="bold", color=NAVY_MPL, pad=14)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.set_facecolor("white"); ax.xaxis.grid(True, linestyle="--", alpha=0.35)
    ax.set_axisbelow(True)
    for bar, val in zip(bars, pvals[::-1]):
        ax.text(bar.get_width()+0.5, bar.get_y()+bar.get_height()/2,
                f"{val:.1f}", va="center", fontsize=9, color="#2C3E50")
    charts.append(_fig_to_bytes(fig))

    # 4 ── Histogram: response rate distribution ──
    rates = [g["response_rate"] for g in groups]
    bins  = [0, 20, 40, 60, 80, 100]
    seg_colors = [RED_MPL, "#FF7043", "#FFA726", "#66BB6A", TEAL_MPL]
    fig, ax = plt.subplots(figsize=(7,4), facecolor="white")
    n, _, patches = ax.hist(rates, bins=bins, edgecolor="white", linewidth=1.5)
    for patch, c in zip(patches, seg_colors):
        patch.set_facecolor(c)
    ax.set_xlabel("Taxa de Resposta (%)", fontsize=11, color="#2C3E50")
    ax.set_ylabel("Nº de Grupos",         fontsize=11, color="#2C3E50")
    ax.set_title("Distribuição da Taxa de Resposta por Grupo", fontsize=13,
                 fontweight="bold", color=NAVY_MPL, pad=14)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.set_facecolor("white"); ax.yaxis.grid(True, linestyle="--", alpha=0.35)
    ax.set_axisbelow(True); ax.set_xticks(bins)
    import math
    max_y = int(max(n)) if len(n) > 0 else 1
    ax.set_yticks(range(0, max_y+2))
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v,_: str(int(v))))
    charts.append(_fig_to_bytes(fig))

    return charts


def make_period_chart(top5_groups, period_m: dict, period_t: dict):
    """
    Figura 5 — Top 5 grupos críticos: chamadas perdidas manhã vs tarde.
    Returns chart bytes.
    """
    names   = [g["label_short"] for g in top5_groups]
    missed_m = [period_m.get(g["group"], {}).get("missed", 0) for g in top5_groups]
    missed_t = [period_t.get(g["group"], {}).get("missed", 0) for g in top5_groups]
    ans_m    = [period_m.get(g["group"], {}).get("answered", 0) for g in top5_groups]
    ans_t    = [period_t.get(g["group"], {}).get("answered", 0) for g in top5_groups]

    y = np.arange(len(names))
    height = 0.32
    fig, axes = plt.subplots(1, 2, figsize=(13, 4.5), facecolor="white")

    # Left subplot — Chamadas Perdidas
    ax = axes[0]
    ax.barh(y + height/2, missed_m, height, label="Manhã",  color=AMBER_MPL, linewidth=0)
    ax.barh(y - height/2, missed_t, height, label="Tarde",  color=NAVY_MPL,  linewidth=0)
    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=9)
    ax.set_xlabel("Chamadas Perdidas", fontsize=10, color="#2C3E50")
    ax.set_title("Chamadas Perdidas", fontsize=12, fontweight="bold", color=NAVY_MPL, pad=10)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.xaxis.grid(True, linestyle="--", alpha=0.35); ax.set_axisbelow(True)
    ax.set_facecolor("white")
    for i, (vm, vt) in enumerate(zip(missed_m, missed_t)):
        if vm: ax.text(vm+0.1, i+height/2, str(vm), va="center", fontsize=8.5, color="#2C3E50")
        if vt: ax.text(vt+0.1, i-height/2, str(vt), va="center", fontsize=8.5, color="#2C3E50")

    # Right subplot — Taxa de Resposta
    ax2 = axes[1]
    rate_m = [period_m.get(g["group"], {}).get("response_rate", 0) for g in top5_groups]
    rate_t = [period_t.get(g["group"], {}).get("response_rate", 0) for g in top5_groups]
    ax2.barh(y + height/2, rate_m, height, label="Manhã", color=AMBER_MPL, linewidth=0)
    ax2.barh(y - height/2, rate_t, height, label="Tarde", color=NAVY_MPL,  linewidth=0)
    ax2.set_yticks(y)
    ax2.set_yticklabels([""] * len(names), fontsize=9)
    ax2.set_xlabel("Taxa de Resposta (%)", fontsize=10, color="#2C3E50")
    ax2.set_title("Taxa de Resposta", fontsize=12, fontweight="bold", color=NAVY_MPL, pad=10)
    ax2.set_xlim(0, 105)
    ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
    ax2.xaxis.grid(True, linestyle="--", alpha=0.35); ax2.set_axisbelow(True)
    ax2.set_facecolor("white")
    for i, (vm, vt) in enumerate(zip(rate_m, rate_t)):
        if vm: ax2.text(vm+0.5, i+height/2, f"{vm:.0f}%", va="center", fontsize=8.5, color="#2C3E50")
        if vt: ax2.text(vt+0.5, i-height/2, f"{vt:.0f}%", va="center", fontsize=8.5, color="#2C3E50")

    patch_m = mpatches.Patch(color=AMBER_MPL, label="Manhã")
    patch_t = mpatches.Patch(color=NAVY_MPL,  label="Tarde")
    fig.legend(handles=[patch_m, patch_t], loc="lower center", ncol=2,
               fontsize=10, framealpha=0, bbox_to_anchor=(0.5, -0.02))

    fig.suptitle("Top 5 Grupos Críticos — Análise Manhã vs. Tarde",
                 fontsize=13, fontweight="bold", color=NAVY_MPL, y=1.02)
    plt.tight_layout()
    return _fig_to_bytes(fig)


def make_quartile_chart(groups):
    """
    Figura de quartis — apenas boxplot horizontal com anotações Q1/Mediana/Q3.
    """
    rates_all = [g["response_rate"] for g in groups]

    q1  = np.percentile(rates_all, 25)
    med = np.percentile(rates_all, 50)
    q3  = np.percentile(rates_all, 75)
    iqr = q3 - q1

    fig, ax = plt.subplots(figsize=(10, 4), facecolor="white")

    ax.boxplot(rates_all, vert=False, patch_artist=True, widths=0.5,
               medianprops={"color": "#" + HEX_RED, "linewidth": 2.8},
               boxprops={"facecolor": "#D6EAF8", "linewidth": 1.5,
                         "edgecolor": NAVY_MPL},
               whiskerprops={"linewidth": 1.5, "linestyle": "--", "color": NAVY_MPL},
               capprops={"linewidth": 2, "color": NAVY_MPL},
               flierprops={"marker": "o", "markerfacecolor": RED_MPL,
                           "markeredgecolor": RED_MPL, "markersize": 6})

    # Jitter strip of individual data points
    jitter = np.random.RandomState(42).uniform(-0.18, 0.18, len(rates_all))
    ax.scatter(rates_all, [1 + j for j in jitter],
               color=TEAL_MPL, alpha=0.55, s=28, zorder=5)

    # Annotations above the box
    for val, lbl, col in [
        (q1,  f"Q1 = {q1:.1f}%",      NAVY_MPL),
        (med, f"Mediana = {med:.1f}%", "#" + HEX_RED),
        (q3,  f"Q3 = {q3:.1f}%",      NAVY_MPL),
    ]:
        ax.annotate(
            lbl, xy=(val, 1), xytext=(val, 1.42),
            fontsize=10, color=col, ha="center", va="bottom",
            fontweight="bold" if "Mediana" in lbl else "normal",
            arrowprops={"arrowstyle": "-", "color": col, "lw": 0.9, "alpha": 0.5},
        )

    # IQR brace below the box
    ax.annotate("", xy=(q3, 0.55), xytext=(q1, 0.55),
                arrowprops={"arrowstyle": "<->", "color": "#7FB3D3", "lw": 2.0})
    ax.text((q1 + q3) / 2, 0.46, f"IQR = {iqr:.1f} pp",
            fontsize=9, color="#5B8EB5", ha="center", va="top", fontstyle="italic")

    ax.set_yticks([])
    ax.set_xlabel("Taxa de Resposta (%)", fontsize=11, color="#2C3E50")
    ax.set_title("Análise de Quartis — Distribuição da Taxa de Resposta",
                 fontsize=13, fontweight="bold", color=NAVY_MPL, pad=14)
    ax.set_xlim(-5, 120)
    ax.set_ylim(0.2, 1.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.set_facecolor("white")
    ax.xaxis.grid(True, linestyle="--", alpha=0.3)
    ax.set_axisbelow(True)

    # Min / Max labels
    mn, mx = min(rates_all), max(rates_all)
    ax.text(mn - 0.5, 0.82, f"Min. {mn:.0f}%", fontsize=8.5, color="#7F8C8D", ha="right", va="center")
    ax.text(mx + 0.5, 0.82, f"Max. {mx:.0f}%", fontsize=8.5, color="#7F8C8D", ha="left",  va="center")

    plt.tight_layout()
    return _fig_to_bytes(fig)


# ──────────────────────────────────────────────
#  AI TEXT GENERATION
# ──────────────────────────────────────────────
VARIANT_STYLES = {
    1: "Técnico e formal — baseado em factos e dados quantitativos precisos. Tom institucional.",
    2: "Analítico-estratégico — identifica padrões, tendências e propõe melhorias operacionais concretas.",
    3: "Executivo e direto — sintetiza os pontos críticos com clareza e apresenta ações prioritárias.",
}


PREMISSAS_DEFAULT = (
    "A portaria/segurança recebeu um maior número de chamadas derivado à escassez de recursos humanos "
    "na área de secretariado devido a saída de funcionários. "
    "As interpretações, introdução e conclusão não devem mencionar nomes de grupos tipo GeralMXXX — apenas assuntos. "
    "Discriminar valores, justificar bem os dados e propor melhorias de forma construtiva e fundamentada."
)


def generate_ai_text(summary, groups, variant: int, api_key: str,
                     period_m=None, period_t=None, premissas: str = ""):
    # Use provided premissas or fall back to the default
    if not premissas or not premissas.strip():
        premissas = PREMISSAS_DEFAULT
    client = anthropic.Anthropic(api_key=api_key)

    top10_priority = sorted(groups, key=lambda x: x["priority_index"], reverse=True)[:10]
    top10_volume   = sorted(groups, key=lambda x: x["presented"],       reverse=True)[:10]
    high_rate = [g for g in groups if g["response_rate"] >= 80]
    low_rate  = [g for g in groups if g["response_rate"] <  40]

    def fmt_group(g):
        nm = g.get("label") or ""
        # Never expose internal group codes in the AI prompt
        if not nm or re.match(r'^GeralM\d+$', nm):
            nm = g.get("servico") or g.get("area") or "Grupo sem designação"
        return (f"{nm}: {g['presented']} recebidas, "
                f"{g['answered']} atendidas, {g['missed']} perdidas, "
                f"taxa {g['response_rate']}%, índice prioridade {g['priority_index']}")

    period_block = ""
    if period_m and period_t:
        top5_codes = [g["group"] for g in top10_priority[:5]]
        lines = []
        for code in top5_codes:
            nm = (period_m.get(code) or period_t.get(code, {})).get("label", code)
            pm = period_m.get(code, {})
            pt = period_t.get(code, {})
            lines.append(
                f"  {nm}: "
                f"manhã — {pm.get('missed',0)} perdidas ({pm.get('response_rate',0):.0f}% atend.); "
                f"tarde — {pt.get('missed',0)} perdidas ({pt.get('response_rate',0):.0f}% atend.)"
            )
        period_block = "\nTOP 5 CRÍTICOS — MANHÃ vs. TARDE:\n" + "\n".join(lines)

    data_block = f"""
PERÍODO: {summary['start_date']} a {summary['end_date']}
TOTAIS GLOBAIS:
  Recebidas: {summary['total_presented']} | Atendidas: {summary['total_answered']} | Perdidas: {summary['total_missed']}
  Taxa global: {summary['pct_answered']*100:.1f}% | Duração total: {summary['total_talking']} | Duração média: {summary['avg_talking']}

TOP 10 POR ÍNDICE DE PRIORIDADE:
{chr(10).join(fmt_group(g) for g in top10_priority)}

TOP 10 POR VOLUME:
{chr(10).join(fmt_group(g) for g in top10_volume)}

Grupos taxa ≥ 80% (bom desempenho): {len(high_rate)} | Grupos taxa < 40% (crítico): {len(low_rate)}
Total de grupos com atividade: {len(groups)}

ANÁLISE DE QUARTIS DA TAXA DE RESPOSTA:
  Q1 (25%): {round(float(__import__("numpy").percentile([g["response_rate"] for g in groups], 25)),1)}%
  Mediana (Q2): {round(float(__import__("numpy").percentile([g["response_rate"] for g in groups], 50)),1)}%
  Q3 (75%): {round(float(__import__("numpy").percentile([g["response_rate"] for g in groups], 75)),1)}%
  IQR: {round(float(__import__("numpy").percentile([g["response_rate"] for g in groups], 75)) - float(__import__("numpy").percentile([g["response_rate"] for g in groups], 25)),1)} pontos percentuais
{period_block}
"""

    has_period = bool(period_m and period_t)
    premissas_block = ""
    if premissas and premissas.strip():
        premissas_block = (
            "\nCONTEXTO E PREMISSAS FORNECIDAS PELO RESPONSÁVEL DO RELATÓRIO:\n"
            + '\"\"\"\n'
            + premissas.strip()
            + '\n\"\"\"\n'
            + "INSTRUÇÕES SOBRE AS PREMISSAS:\n"
            + "- Integra este contexto de forma natural e profissional no texto — não o cites literalmente.\n"
            + "- Usa estas informações para justificar os dados apresentados, explicar desvios e contextualizar o desempenho.\n"
            + "- Onde existam constrangimentos externos (ex: saída de recursos humanos), menciona-os com objetividade e sem dramatismo.\n"
            + "- Propõe melhorias concretas e construtivas diretamente relacionadas com o contexto fornecido.\n"
            + "- Mantém sempre uma imagem clara e eficiente da instituição.\n"
        )
    prompt = f"""És um redator técnico especializado em relatórios institucionais de atendimento telefónico para organismos públicos portugueses. Escreve em português europeu correto, fluido e profissional.

ESTILO DE REDAÇÃO: {VARIANT_STYLES[variant]}

DADOS ESTATÍSTICOS DO PERÍODO:
{data_block}
{premissas_block}
{"NOTA: Na secção 'analise', dedica um parágrafo específico às diferenças de desempenho entre manhã e tarde nos grupos críticos." if has_period else ""}

INSTRUÇÕES GERAIS:
- CRÍTICO: NUNCA uses códigos de grupo como "GeralM509", "GeralM909", "GeralM219" ou similares. Refere SEMPRE os grupos pelo seu assunto/nome completo.
- Cada secção deve ser coesa, bem estruturada e com argumentação clara.
- Usa os dados numéricos para fundamentar as afirmações — não faças afirmações vagas.
- O tom deve ser sempre construtivo, institucional e orientado para a melhoria.
- Não uses linguagem excessivamente técnica nem siglas sem explicação prévia.
- As conclusões devem incluir pelo menos 2-3 recomendações práticas e específicas.
- Inclui informação dos quartis da taxa de resposta para enquadrar o desempenho relativo dos grupos.

Devolve APENAS JSON válido (sem markdown, sem texto adicional):
{{
  "introducao": "<3-4 parágrafos sep. por \\n\\n. Âmbito do relatório, período em análise, enquadramento institucional, síntese dos resultados globais e principais condicionantes do período.>",
  "analise":    "<5-7 parágrafos sep. por \\n\\n. Análise detalhada: desempenho global, grupos críticos com justificação dos dados, grupos de bom desempenho, padrões identificados{', diferenças de período manhã/tarde' if has_period else ''}. Fundamenta cada afirmação com números.>",
  "conclusoes": "<3-4 parágrafos sep. por \\n\\n. Síntese dos pontos críticos, recomendações específicas e acionáveis, perspetivas de melhoria e próximos passos sugeridos.>"
}}"""

    msg = client.messages.create(
        model=_CFG_MODEL,
        max_tokens=_CFG_TOKENS,
        messages=[{"role":"user","content":prompt}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```json\s*","",raw); raw = re.sub(r"\s*```$","",raw)
    return json.loads(raw)


# ──────────────────────────────────────────────
#  DOCX HELPERS
# ──────────────────────────────────────────────
def _cell_shading(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)


def _para_border(para, side="bottom", color=HEX_NAVY, size=8):
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None: pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bd   = OxmlElement(f"w:{side}")
    bd.set(qn("w:val"),"single"); bd.set(qn("w:sz"),str(size))
    bd.set(qn("w:space"),"1");    bd.set(qn("w:color"),color)
    pBdr.append(bd)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None: pStyle.addnext(pBdr)
    else: pPr.insert(0, pBdr)


def _add_run(para, text, bold=False, italic=False,
             size=10.5, color=None, font="Calibri"):
    r = para.add_run(text)
    r.font.name = font; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return r


def _section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(8)
    _add_run(p, f"{number}. ", bold=True, size=14, color=C_BLUE)
    _add_run(p, title,         bold=True, size=14, color=C_NAVY)
    _para_border(p, "bottom", HEX_BLUE, 6)
    return p


def _body_paragraphs(doc, text: str):
    for chunk in text.strip().split("\n\n"):
        chunk = chunk.strip()
        if not chunk: continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after  = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        _add_run(p, chunk, size=10.5, color=C_TEXT)


def _set_repeat_header(row):
    """Mark a table row to repeat as header on every page."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    tblH = OxmlElement("w:tblHeader")
    tblH.set(qn("w:val"), "true")
    trPr.append(tblH)


def _cant_split(row):
    """Prevent a table row from being split across a page break."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    cs   = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "true")
    trPr.append(cs)


def _table_header_row(table, headers, col_widths_cm, bg=HEX_NAVY):
    row = table.rows[0]
    for ci, (h, w) in enumerate(zip(headers, col_widths_cm)):
        cell = row.cells[ci]
        cell.width = Cm(w)
        _cell_shading(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, h, bold=True, size=9, color=C_WHITE)
    _set_repeat_header(row)   # repeat on every page


def _set_header_footer(doc, section):
    """Apply template header/footer images to the document section."""
    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)
    hp = header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)
    run = hp.add_run()
    if os.path.exists(HEADER_IMG):
        run.add_picture(HEADER_IMG, width=Cm(16.8))

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    run2 = fp.add_run()
    if os.path.exists(FOOTER_IMG):
        run2.add_picture(FOOTER_IMG, width=Cm(16.8))


# ──────────────────────────────────────────────
#  WORD REPORT GENERATION
# ──────────────────────────────────────────────
def build_word_report(summary, groups, ai_text: dict, charts: list, variant_num: int,
                      period_m=None, period_t=None, chart_period=None,
                      chart_quartile=None) -> bytes:
    doc = Document()

    # ── Page geometry (A4, match template margins) ──
    sec = doc.sections[0]
    sec.page_width        = Cm(21)
    sec.page_height       = Cm(29.7)
    sec.left_margin       = Cm(1.7)
    sec.right_margin      = Cm(1.7)
    sec.top_margin        = Cm(4.5)   # room for header image
    sec.bottom_margin     = Cm(2.0)
    sec.header_distance   = Cm(1.2)
    sec.footer_distance   = Cm(0.8)

    # ── Default style ──
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = C_TEXT

    # ── Template header/footer ──
    _set_header_footer(doc, sec)

    # ══════════════════════════════════
    #  CAPA
    # ══════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "RELATÓRIO DE ATENDIMENTO", bold=True, size=26, color=C_NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    _add_run(p2, "Análise de Desempenho e Reencaminhamento de Chamadas",
             italic=True, size=13, color=C_BLUE)
    _para_border(p2, "bottom", HEX_BLUE, 6)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(28)
    _add_run(p3, f"Período de Referência: {summary['start_date']} a {summary['end_date']}",
             size=11, color=C_GREY)

    # KPI strip
    kpi_items = [
        ("Chamadas\nRecebidas",  str(summary["total_presented"]), HEX_NAVY),
        ("Chamadas\nAtendidas",  str(summary["total_answered"]),  HEX_TEAL),
        ("Chamadas\nPerdidas",   str(summary["total_missed"]),    HEX_RED),
        ("Taxa Global\nResposta", f"{summary['pct_answered']*100:.1f}%",
         HEX_GREEN if summary["pct_answered"] >= 0.6 else HEX_ORANGE),
    ]
    kpi_tbl = doc.add_table(rows=2, cols=4)
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (lbl, val, color) in enumerate(kpi_items):
        vc = kpi_tbl.rows[0].cells[ci]; vc.width = Cm(3.8)
        _cell_shading(vc, color)
        vp = vc.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(vp, val, bold=True, size=22, color=C_WHITE)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lc = kpi_tbl.rows[1].cells[ci]; lc.width = Cm(3.8)
        _cell_shading(lc, color)
        lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(lp, lbl, size=8.5, color=RGBColor(0xCC, 0xE5, 0xFF))
    doc.add_paragraph()

    # ══════════════════════════════════
    #  SECÇÃO 1 — INTRODUÇÃO
    # ══════════════════════════════════
    doc.add_page_break()
    _section_heading(doc, 1, "Introdução")
    _body_paragraphs(doc, ai_text["introducao"])

    # ══════════════════════════════════
    #  SECÇÃO 2 — RESUMO ESTATÍSTICO
    # ══════════════════════════════════
    _section_heading(doc, 2, "Resumo Estatístico")
    summ_data = [
        ("Total de Chamadas Recebidas",  str(summary["total_presented"])),
        ("Total de Chamadas Atendidas",  str(summary["total_answered"])),
        ("Total de Chamadas Perdidas",   str(summary["total_missed"])),
        ("Taxa Global de Resposta",       f"{summary['pct_answered']*100:.1f}%"),
        ("Duração Total de Atendimento", summary["total_talking"]),
        ("Duração Média de Atendimento", summary["avg_talking"]),
    ]
    tbl = doc.add_table(rows=len(summ_data)+1, cols=2)
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_header_row(tbl, ["Indicador", "Valor"], [10, 5.5])
    for ri, (lbl2, val) in enumerate(summ_data):
        bg = HEX_LGREY if ri%2==0 else HEX_WHITE
        cells = tbl.rows[ri+1].cells
        cells[0].width = Cm(10); cells[1].width = Cm(5.5)
        _cell_shading(cells[0], bg); _cell_shading(cells[1], bg)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cells[0].paragraphs[0], lbl2, size=10, color=C_TEXT)
        _add_run(cells[1].paragraphs[0], val,  size=10, bold=True, color=C_NAVY)
    doc.add_paragraph()

    # ══════════════════════════════════
    #  SECÇÃO 3 — ANÁLISE POR GRUPO
    # ══════════════════════════════════
    _section_heading(doc, 3, "Análise por Grupo")
    p_note = doc.add_paragraph()
    _add_run(p_note,
             "Grupos ordenados por Índice de Prioridade (descendente). "
             "A vermelho os grupos com índice crítico (> 20).",
             italic=True, size=9.5, color=C_GREY)
    p_note.paragraph_format.space_after = Pt(8)

    col_hdrs = ["Assunto / Tema", "Chamadas Recebidas", "Chamadas Atendidas", "Chamadas Perdidas", "Taxa de Resposta", "Índice de Prioridade"]
    col_ws   = [7.5, 2.0, 2.0, 2.0, 2.0, 2.2]
    dtbl = doc.add_table(rows=len(groups)+1, cols=6)
    dtbl.style = "Table Grid"; dtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_header_row(dtbl, col_hdrs, col_ws)
    for ri, g in enumerate(groups):
        bg  = HEX_LGREY if ri%2==0 else HEX_WHITE
        row = dtbl.rows[ri+1]
        _cant_split(row)   # never break a data row across a page
        vals = [g.get("label") or g["group"],
                str(g["presented"]), str(g["answered"]), str(g["missed"]),
                f"{g['response_rate']:.1f}%", f"{g['priority_index']:.1f}"]
        for ci, (val, w) in enumerate(zip(vals, col_ws)):
            cell = row.cells[ci]; cell.width = Cm(w)
            critical = g["priority_index"] > 20
            cell_bg  = "FFF0F0" if (critical and ci==5) else bg
            _cell_shading(cell, cell_bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
            color = C_RED if (critical and ci==5) else C_TEXT
            _add_run(p, val, bold=(critical and ci==5), size=8.5 if ci==0 else 9, color=color)

    doc.add_paragraph()
    p_pi = doc.add_paragraph()
    p_pi.paragraph_format.left_indent  = Cm(0.5)
    p_pi.paragraph_format.space_before = Pt(4)
    p_pi.paragraph_format.space_after  = Pt(8)
    _para_border(p_pi, "left", HEX_BLUE, 18)
    _add_run(p_pi, "Índice de Prioridade", bold=True, size=10, color=C_NAVY)
    _add_run(p_pi, "  ·  Fórmula: Chamadas Perdidas × (1 − Taxa de Resposta)",
             italic=True, size=9.5, color=C_GREY)

    # ══════════════════════════════════
    #  SECÇÃO 4 — GRÁFICOS (existentes)
    # ══════════════════════════════════
    _section_heading(doc, 4, "Gráficos e Visualizações")
    chart_meta = [
        ("Figura 1 — Distribuição Global de Chamadas",            Inches(3.6)),
        ("Figura 2 — Top 15 Grupos: Atendidas vs. Perdidas",      Inches(6.2)),
        ("Figura 3 — Top 10 Grupos por Índice de Prioridade",     Inches(5.5)),
        ("Figura 4 — Distribuição da Taxa de Resposta por Grupo", Inches(5.0)),
    ]
    for chart_bytes, (caption, width) in zip(charts, chart_meta):
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(10); pc.paragraph_format.space_after = Pt(4)
        _add_run(pc, caption, italic=True, size=9, color=C_GREY)
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_after = Pt(14)
        pi.add_run().add_picture(io.BytesIO(chart_bytes), width=width)

    # Figura 5 — Análise de Quartis
    if chart_quartile:
        pc_q = doc.add_paragraph()
        pc_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc_q.paragraph_format.space_before = Pt(10); pc_q.paragraph_format.space_after = Pt(4)
        _add_run(pc_q, "Figura 5 — Análise de Quartis da Taxa de Resposta por Grupo",
                 italic=True, size=9, color=C_GREY)
        pi_q = doc.add_paragraph()
        pi_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi_q.paragraph_format.space_after = Pt(14)
        pi_q.add_run().add_picture(io.BytesIO(chart_quartile), width=Inches(6.3))

        # Box with quartile legend
        p_ql = doc.add_paragraph()
        p_ql.paragraph_format.left_indent  = Cm(0.5)
        p_ql.paragraph_format.space_before = Pt(2)
        p_ql.paragraph_format.space_after  = Pt(10)
        _para_border(p_ql, "left", HEX_TEAL, 18)
        _add_run(p_ql, "Leitura do gráfico  ", bold=True, size=9.5, color=C_NAVY)
        _add_run(p_ql,
                 "O boxplot mostra a mediana, os quartis Q1 e Q3, a amplitude interquartil (IQR) "
                 "e a distribuição individual de todos os grupos (pontos sobrepostos).",
                 italic=True, size=9, color=C_GREY)

    # ══════════════════════════════════
    #  SECÇÃO 5 — ANÁLISE MANHÃ vs TARDE  (opcional)
    # ══════════════════════════════════
    sec_num = 5
    if period_m and period_t and chart_period:
        _section_heading(doc, sec_num, "Análise por Período — Manhã vs. Tarde")
        sec_num += 1

        p_intro = doc.add_paragraph()
        _add_run(p_intro,
                 "A tabela e gráfico seguintes detalham os cinco grupos com maior Índice de Prioridade, "
                 "discriminando o volume de chamadas recebidas, perdidas e a taxa de resposta por período do dia.",
                 size=10.5, color=C_TEXT)
        p_intro.paragraph_format.space_after = Pt(10)

        top5 = groups[:5]

        # Table: top 5 × manhã/tarde
        pt_hdrs = ["Assunto / Tema",
                   "Manhã — Recebidas", "Manhã — Perdidas", "Manhã — Taxa de Resposta",
                   "Tarde — Recebidas", "Tarde — Perdidas", "Tarde — Taxa de Resposta"]
        pt_ws   = [5.2, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7]

        ptbl = doc.add_table(rows=len(top5)+1, cols=7)
        ptbl.style = "Table Grid"; ptbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Bicolor header: manhã columns amber, tarde columns navy
        hrow = ptbl.rows[0]
        for ci, (h, w) in enumerate(zip(pt_hdrs, pt_ws)):
            cell = hrow.cells[ci]; cell.width = Cm(w)
            bg = HEX_NAVY if ci == 0 else (HEX_ORANGE if ci <= 3 else HEX_NAVY)
            _cell_shading(cell, bg)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, h, bold=True, size=8.5, color=C_WHITE)

        for ri, g in enumerate(top5):
            bg  = HEX_LGREY if ri%2==0 else HEX_WHITE
            row = ptbl.rows[ri+1]
            pm  = period_m.get(g["group"], {})
            pt  = period_t.get(g["group"], {})
            vals = [
                g.get("label") or g["group"],
                str(pm.get("presented", 0)), str(pm.get("missed", 0)),
                f"{pm.get('response_rate', 0):.0f}%",
                str(pt.get("presented", 0)), str(pt.get("missed", 0)),
                f"{pt.get('response_rate', 0):.0f}%",
            ]
            for ci, (val, w) in enumerate(zip(vals, pt_ws)):
                cell = row.cells[ci]; cell.width = Cm(w)
                _cell_shading(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, val, size=8.5 if ci==0 else 9, color=C_TEXT)

        doc.add_paragraph()
        pc5 = doc.add_paragraph()
        pc5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc5.paragraph_format.space_before = Pt(6); pc5.paragraph_format.space_after = Pt(4)
        _add_run(pc5, "Figura 5 — Top 5 Grupos Críticos: Chamadas Perdidas e Taxa de Resposta por Período",
                 italic=True, size=9, color=C_GREY)
        pi5 = doc.add_paragraph()
        pi5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi5.paragraph_format.space_after = Pt(14)
        pi5.add_run().add_picture(io.BytesIO(chart_period), width=Inches(6.3))

    # ══════════════════════════════════
    #  SECÇÃO N — ANÁLISE
    # ══════════════════════════════════
    _section_heading(doc, sec_num, "Análise e Interpretação dos Dados")
    _body_paragraphs(doc, ai_text["analise"])
    sec_num += 1

    # ══════════════════════════════════
    #  SECÇÃO N+1 — CONCLUSÕES
    # ══════════════════════════════════
    _section_heading(doc, sec_num, "Conclusões e Recomendações")
    _body_paragraphs(doc, ai_text["conclusoes"])

    # ── Serialize + fix python-docx zoom bug ──
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    import zipfile as _zf
    fixed = io.BytesIO()
    with _zf.ZipFile(io.BytesIO(raw), "r") as zin, \
         _zf.ZipFile(fixed, "w", _zf.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                txt  = data.decode("utf-8")
                txt  = txt.replace('<w:zoom w:val="bestFit"/>',
                                   '<w:zoom w:val="bestFit" w:percent="100"/>')
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    return fixed.getvalue()


# ──────────────────────────────────────────────
#  STREAMLIT UI
# ──────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="Gerador de Relatório de Atendimento",
        page_icon="📞", layout="centered",
    )
    st.markdown("""
    <style>
    .block-container { max-width: 860px; padding-top: 2rem; }
    h1 { color: #1B3A5C; letter-spacing: -0.5px; }
    h3 { color: #4A90D9; }
    .stDownloadButton > button {
        background: #1B3A5C; color: white; border: none;
        border-radius: 6px; font-weight: 600; padding: 0.55rem 1.6rem;
    }
    .stDownloadButton > button:hover { background: #4A90D9; }
    .stButton > button {
        background: #1B3A5C; color: white; border: none;
        border-radius: 6px; font-weight: 600; padding: 0.55rem 1.6rem;
    }
    .stButton > button:hover { background: #4A90D9; }
    div[data-testid="metric-container"] {
        background: #F0F5FA; border-radius: 8px; padding: 1rem; border: 1px solid #D6EAF8;
    }
    </style>
    """, unsafe_allow_html=True)

    st.title("📞 Gerador de Relatório de Atendimento")
    st.markdown(
        "Carregue o ficheiro Excel de dados globais (e opcionalmente os dados "
        "de manhã e de tarde) para gerar um relatório profissional Word com análise por IA."
    )

    # ── API Key — lida de config.py, env var, ou campo manual ──
    api_key = _CFG_KEY or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        with st.expander("🔑 Configuração da API Key", expanded=True):
            api_key = st.text_input(
                "Anthropic API Key",
                type="password",
                placeholder="sk-ant-...",
                help="Pode pré-configurar em config.py para não precisar introduzir manualmente.",
            )
            st.caption("Para pré-configurar permanentemente, edite `config.py` e defina `ANTHROPIC_API_KEY`.")

    st.divider()

    # ── Modo de input ──
    st.markdown("### 📂 Ficheiros de Dados")
    mode = st.radio(
        "Modo de análise:",
        ["1 ficheiro (dados globais)", "3 ficheiros (global + manhã + tarde)"],
        horizontal=True,
        help="3 ficheiros ativa a análise de período nos 5 grupos mais críticos.",
    )
    three_mode = "3 ficheiros" in mode

    col_main, col_m, col_t = st.columns([2, 1, 1]) if three_mode else [st.container(), None, None]

    with col_main:
        uploaded_main = st.file_uploader(
            "Ficheiro global (.xlsx)", type=["xlsx"],
            help="Exportação completa do sistema de centralita.",
        )

    uploaded_m = uploaded_t = None
    if three_mode:
        with col_m:
            uploaded_m = st.file_uploader("Manhã (.xlsx)", type=["xlsx"])
        with col_t:
            uploaded_t = st.file_uploader("Tarde (.xlsx)", type=["xlsx"])

    if not uploaded_main:
        st.info("👆 Carregue pelo menos o ficheiro de dados globais para começar.")
        return

    # ── Parse ──
    with st.spinner("A ler os ficheiros…"):
        try:
            summary, groups = parse_excel(uploaded_main)
        except Exception as e:
            st.error(f"Erro ao ler o ficheiro global: {e}"); return

        period_m = period_t = None
        if three_mode and uploaded_m and uploaded_t:
            try:
                period_m = parse_excel_period(uploaded_m, "Manhã")
                period_t = parse_excel_period(uploaded_t, "Tarde")
            except Exception as e:
                st.warning(f"Erro nos ficheiros de período: {e}")
                period_m = period_t = None

    st.success(
        f"✅ **{summary['total_presented']}** chamadas  |  "
        f"Período: **{summary['start_date']}** a **{summary['end_date']}**"
        + (f"  |  📊 Dados manhã/tarde carregados" if period_m else "")
    )

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Recebidas",     summary["total_presented"])
    c2.metric("Atendidas",     summary["total_answered"])
    c3.metric("Perdidas",      summary["total_missed"])
    c4.metric("Taxa Resposta", f"{summary['pct_answered']*100:.1f}%")

    st.divider()

    # ── Premissas / Contexto ──
    st.markdown("### 📝 Premissas e Contexto do Relatório")
    st.markdown(
        "Introduza aqui o contexto específico do período em análise — "
        "condicionantes, explicações para desvios, observações relevantes. "
        "A IA usará esta informação para fundamentar e contextualizar a narrativa do relatório."
    )
    premissas = st.text_area(
        "Premissas / Contexto (opcional)",
        height=180,
        placeholder=(
            "Exemplo:\n"
            "- A linha de apoio recebeu um volume superior de chamadas devido à saída de recursos humanos "
            "de alguns secretariados.\n"
            "- O grupo X registou menor taxa de resposta por razões de reestruturação interna.\n"
            "- Propor medidas para reforço do atendimento nos períodos da tarde."
        ),
        help="O texto aqui inserido é usado pela IA para contextualizar a análise e as conclusões. "
             "Seja específico e objetivo — mais contexto = análise mais fundamentada.",
        label_visibility="collapsed",
    )
    if premissas.strip():
        st.success(f"✏️ Contexto registado ({len(premissas.split())} palavras) — será incorporado na narrativa.")

    st.divider()

    # ── Opções ──
    st.markdown("### ⚙️ Opções de Geração")
    n_variants = st.radio(
        "Número de hipóteses de relatório:",
        options=[1, 2, 3], horizontal=True,
        help="Cada hipótese tem um estilo de redação diferente.",
    )
    variant_labels = {1:"Técnico e Formal", 2:"Analítico-Estratégico", 3:"Executivo e Conciso"}
    if n_variants > 1:
        st.caption(" · ".join(f"H{i}: **{variant_labels[i]}**" for i in range(1, n_variants+1)))

    if st.button("🚀 Gerar Relatório" + ("s" if n_variants > 1 else "")):
        if not api_key:
            st.error("Por favor insira a Anthropic API Key antes de continuar."); return

        with st.spinner("A gerar gráficos…"):
            try:
                charts = make_charts(summary, groups)
                chart_quartile = make_quartile_chart(groups)
                chart_period = None
                if period_m and period_t:
                    chart_period = make_period_chart(groups[:5], period_m, period_t)
            except Exception as e:
                st.error(f"Erro ao gerar gráficos: {e}"); return

        results = []
        for i in range(1, n_variants+1):
            lbl = variant_labels[i]
            with st.spinner(f"A gerar Hipótese {i} ({lbl}) com IA…"):
                try:
                    ai_text   = generate_ai_text(summary, groups, i, api_key, period_m, period_t, premissas)
                    doc_bytes = build_word_report(summary, groups, ai_text, charts, i,
                                                  period_m, period_t, chart_period,
                                                  chart_quartile)
                    results.append((i, lbl, doc_bytes))
                except Exception as e:
                    st.error(f"Erro na hipótese {i}: {e}"); return

        st.success("✅ Relatório(s) gerado(s) com sucesso!")
        st.divider()
        st.markdown("### ⬇️ Descarregar")
        period_str = f"{summary['start_date'].replace('/','')}-{summary['end_date'].replace('/','')}"
        for i, lbl, doc_bytes in results:
            fname = f"Relatorio_Atendimento_{period_str}_H{i}_{lbl.replace(' ','_')}.docx"
            st.download_button(
                label=f"📄 Hipótese {i} — {lbl}", data=doc_bytes,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{i}",
            )
        if n_variants > 1:
            st.caption("Descarregue todas as hipóteses e escolha a que melhor se adapta.")


if __name__ == "__main__":
    main()
